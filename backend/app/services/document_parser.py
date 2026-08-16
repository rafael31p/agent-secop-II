"""Extracción de texto de PDF, DOCX y TXT subidos por el usuario."""

from __future__ import annotations

import io
import logging

log = logging.getLogger(__name__)

LIMITE_CARACTERES = 400_000
EXTENSIONES = {".pdf", ".docx", ".txt", ".md"}


class DocumentoNoSoportado(ValueError):
    pass


def extraer_texto(
    nombre_archivo: str, contenido: bytes
) -> tuple[str, str, int | None, bool]:
    """Devuelve `(texto, tipo, paginas, truncado)`."""
    nombre = (nombre_archivo or "").lower()

    if nombre.endswith(".pdf"):
        texto, paginas = _pdf(contenido)
        tipo = "pdf"
    elif nombre.endswith(".docx"):
        texto, paginas = _docx(contenido), None
        tipo = "docx"
    elif nombre.endswith((".txt", ".md")):
        texto, paginas = contenido.decode("utf-8", errors="replace"), None
        tipo = "texto"
    else:
        raise DocumentoNoSoportado(
            f"Formato no soportado. Usa: {', '.join(sorted(EXTENSIONES))}. "
            "Los .doc antiguos deben convertirse a .docx o PDF."
        )

    texto = _limpiar(texto)
    truncado = len(texto) > LIMITE_CARACTERES
    if truncado:
        texto = texto[:LIMITE_CARACTERES]
    return texto, tipo, paginas, truncado


def _pdf(contenido: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    lector = PdfReader(io.BytesIO(contenido))
    partes: list[str] = []
    for indice, pagina in enumerate(lector.pages, start=1):
        try:
            texto = pagina.extract_text() or ""
        except Exception as exc:  # una página corrupta no debe tumbar el documento
            log.warning("No se pudo extraer la página %s: %s", indice, exc)
            texto = ""
        partes.append(f"\n\n--- Página {indice} ---\n{texto}")
    return "".join(partes), len(lector.pages)


def _docx(contenido: bytes) -> str:
    import docx

    documento = docx.Document(io.BytesIO(contenido))
    partes = [p.text for p in documento.paragraphs if p.text.strip()]

    # Las tablas de los anexos técnicos suelen contener los requisitos.
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))
    return "\n".join(partes)


def _limpiar(texto: str) -> str:
    lineas = [linea.rstrip() for linea in texto.replace("\r\n", "\n").split("\n")]
    salida: list[str] = []
    vacias = 0
    for linea in lineas:
        if linea:
            vacias = 0
            salida.append(linea)
        else:
            vacias += 1
            if vacias <= 2:
                salida.append("")
    return "\n".join(salida).strip()
